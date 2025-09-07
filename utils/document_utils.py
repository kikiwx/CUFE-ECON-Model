"""
文档处理工具
中央财经大学经济学院 - 经济学大模型聊天助手
"""

import re
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor

logger = logging.getLogger(__name__)


def clean_text_for_word(text):
    """清理文本内容，移除特殊字符和格式"""
    if not text:
        return ""

    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 粗体
    text = re.sub(r'\*(.*?)\*', r'\1', text)  # 斜体
    text = re.sub(r'#{1,6}\s*', '', text)  # 标题符号
    text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)  # 代码块
    text = re.sub(r'`([^`]*)`', r'\1', text)  # 行内代码

    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)

    return text.strip()


def set_font_style(run, font_name="SimSun", font_size=12, bold=False, color=None):
    """设置字体样式"""
    try:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold

        if color:
            run.font.color.rgb = color

        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception as e:
        logger.warning(f"设置字体样式失败: {e}")


def add_table_with_style(doc, rows, cols, data=None):
    """添加带样式的表格"""
    try:
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        if data:
            for i, row_data in enumerate(data):
                if i < rows:
                    for j, cell_data in enumerate(row_data):
                        if j < cols:
                            cell = table.cell(i, j)
                            cell.text = str(cell_data)

                            if i == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        set_font_style(run, "SimSun", 11, bold=True)
                            else:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        set_font_style(run, "SimSun", 11)

        return table
    except Exception as e:
        logger.error(f"添加表格失败: {e}")
        return None


def create_word_document(report_data):
    """创建Word文档"""
    try:
        doc = Document()

        if not report_data.get('outline'):
            raise ValueError("报告大纲数据缺失")

        outline = report_data['outline']
        sections = report_data.get('sections', {})

        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = Pt(12)

        if outline.get('title'):
            title = doc.add_heading(outline['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                set_font_style(run, "SimSun", 18, bold=True)
        else:
            title = doc.add_heading('经济学报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                set_font_style(run, "SimSun", 18, bold=True)

        info_para = doc.add_paragraph()
        info_run = info_para.add_run(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
        set_font_style(info_run, "SimSun", 10)
        info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        info_data = [
            ('报告主题', report_data.get('topic', '未指定')),
            ('章节数量', str(len(sections))),
            ('生成系统', 'CUFE经济学大模型')
        ]

        info_table = add_table_with_style(doc, 3, 2, info_data)
        doc.add_paragraph()

        if outline.get('abstract'):
            abstract_heading = doc.add_heading('摘要', level=1)
            for run in abstract_heading.runs:
                set_font_style(run, "SimSun", 16, bold=True)

            abstract_content = clean_text_for_word(outline['abstract'])
            abstract_paragraph = doc.add_paragraph()
            abstract_run = abstract_paragraph.add_run(abstract_content)
            set_font_style(abstract_run, "SimSun", 12)

        if outline.get('sections'):
            toc_heading = doc.add_heading('目录', level=1)
            for run in toc_heading.runs:
                set_font_style(run, "SimSun", 16, bold=True)

            for section in outline['sections']:
                toc_paragraph = doc.add_paragraph(f"{section.get('id', '')}. {section.get('title', '')}")
                for run in toc_paragraph.runs:
                    set_font_style(run, "SimSun", 12)

            doc.add_page_break()


        if outline.get('sections'):
            for section in outline['sections']:
                section_id = section.get('id')
                section_title = section.get('title', '')

                if section_title:
                    section_heading = doc.add_heading(f"{section_id}. {section_title}", level=1)
                    for run in section_heading.runs:
                        set_font_style(run, "SimSun", 16, bold=True)

                if section_id in sections and sections[section_id].get('content'):
                    content = clean_text_for_word(sections[section_id]['content'])

                    paragraphs = content.split('\n\n')
                    for para_text in paragraphs:
                        para_text = para_text.strip()
                        if para_text:
                            paragraph = doc.add_paragraph()

                            if '「' in para_text and '」' in para_text:
                                parts = re.split(r'[「」]', para_text)
                                for i, part in enumerate(parts):
                                    if part:
                                        run = paragraph.add_run(part)

                                        set_font_style(run, "SimSun", 12, bold=(i % 2 == 1))
                            else:
                                para_run = paragraph.add_run(para_text)
                                set_font_style(para_run, "SimSun", 12)
                else:

                    placeholder_paragraph = doc.add_paragraph()
                    placeholder_run = placeholder_paragraph.add_run("本章节内容正在完善中...")
                    set_font_style(placeholder_run, "SimSun", 12)


        doc.add_page_break()
        footer_heading = doc.add_heading('生成信息', level=1)
        for run in footer_heading.runs:
            set_font_style(run, "SimSun", 16, bold=True)

        footer_paragraph = doc.add_paragraph()
        footer_text = f"""本报告由CUFE经济学大模型自动生成
生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}
报告ID：{report_data.get('id', 'N/A')}
中央财经大学经济学院出品"""

        footer_run = footer_paragraph.add_run(footer_text)
        set_font_style(footer_run, "SimSun", 10)
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        logger.info("Word文档创建完成")
        return doc

    except Exception as e:
        logger.error(f"创建Word文档失败: {str(e)}", exc_info=True)
        doc = Document()
        error_title = doc.add_heading('报告生成错误', 0)
        error_para = doc.add_paragraph(f"报告生成时遇到错误：{str(e)}")
        return doc


def create_markdown_document(report_data):
    """创建Markdown文档"""
    try:
        if not report_data.get('outline'):
            raise ValueError("报告大纲数据缺失")

        outline = report_data['outline']
        sections = report_data.get('sections', {})

        markdown_content = []

        if outline.get('title'):
            markdown_content.append(f"# {outline['title']}\n")

        markdown_content.append(f"**生成时间**: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}\n")
        markdown_content.append(f"**生成系统**: CUFE经济学大模型\n")
        markdown_content.append(f"**报告主题**: {report_data.get('topic', '未指定')}\n\n")

        if outline.get('abstract'):
            markdown_content.append("## 摘要\n\n")
            markdown_content.append(f"{outline['abstract']}\n\n")

        if outline.get('sections'):
            markdown_content.append("## 目录\n\n")
            for section in outline['sections']:
                markdown_content.append(f"- {section.get('id', '')}. {section.get('title', '')}\n")
            markdown_content.append("\n---\n\n")

        if outline.get('sections'):
            for section in outline['sections']:
                section_id = section.get('id')
                section_title = section.get('title', '')

                if section_title:
                    markdown_content.append(f"## {section_id}. {section_title}\n\n")

                if section_id in sections and sections[section_id].get('content'):
                    content = sections[section_id]['content']
                    markdown_content.append(f"{content}\n\n")
                else:
                    markdown_content.append("本章节内容正在完善中...\n\n")


        markdown_content.append("---\n\n")
        markdown_content.append("## 生成信息\n\n")
        markdown_content.append(f"- **生成时间**: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}\n")
        markdown_content.append(f"- **报告ID**: {report_data.get('id', 'N/A')}\n")
        markdown_content.append("- **出品方**: 中央财经大学经济学院\n")

        return ''.join(markdown_content)

    except Exception as e:
        logger.error(f"创建Markdown文档失败: {str(e)}", exc_info=True)
        return f"# 报告生成错误\n\n报告生成时遇到错误：{str(e)}"



def validate_document_data(report_data):
    """验证文档数据的完整性"""
    errors = []

    if not report_data:
        errors.append("报告数据为空")
        return errors

    if not report_data.get('outline'):
        errors.append("缺少报告大纲")
    else:
        outline = report_data['outline']
        if not outline.get('title'):
            errors.append("缺少报告标题")
        if not outline.get('sections'):
            errors.append("缺少章节信息")

    if not report_data.get('sections'):
        errors.append("缺少章节内容")

    return errors


def get_document_statistics(report_data):
    """获取文档统计信息"""
    stats = {
        'title_length': 0,
        'abstract_length': 0,
        'total_sections': 0,
        'total_words': 0,
        'completed_sections': 0
    }

    try:
        outline = report_data.get('outline', {})
        sections = report_data.get('sections', {})

        if outline.get('title'):
            stats['title_length'] = len(outline['title'])

        if outline.get('abstract'):
            stats['abstract_length'] = len(outline['abstract'])
            stats['total_words'] += len(outline['abstract'])

        if outline.get('sections'):
            stats['total_sections'] = len(outline['sections'])

        for section_data in sections.values():
            if section_data.get('content'):
                stats['completed_sections'] += 1
                stats['total_words'] += len(section_data['content'])

    except Exception as e:
        logger.error(f"获取文档统计失败: {e}")

    return stats